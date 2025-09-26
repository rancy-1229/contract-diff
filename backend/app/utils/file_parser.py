import fitz  # PyMuPDF
import json
import os
from typing import Dict, List
from app.config import settings

class DocumentParser:
    def __init__(self):
        self.upload_dir = settings.DOCUMENTS_DIR
        self.image_dir = settings.IMAGES_DIR
        self.temp_dir = settings.TEMP_DIR
        
        # 确保目录存在
        for dir_path in [self.upload_dir, self.image_dir, self.temp_dir]:
            os.makedirs(dir_path, exist_ok=True)
    
    async def parse_document(self, file_path: str, file_type: str) -> Dict:
        """解析文档并返回结构化数据 - 按照5步流程实现"""
        print(f"[DEBUG] ===== 开始文档解析流程 =====")
        print(f"[DEBUG] 文档路径: {file_path}")
        print(f"[DEBUG] 文档类型: {file_type}")
        
        # 步骤1: 格式转换 (docx -> pdf)
        pdf_path = None
        if file_type.lower() == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            print("[DEBUG] 步骤1: Word文档转换为PDF")
            pdf_path = await self._convert_docx_to_pdf(file_path)
            if not pdf_path:
                print("[DEBUG] 转换失败，使用简化解析")
                return await self._parse_word_simple(file_path)
            file_path = pdf_path
        elif file_type.lower() == 'application/pdf':
            print("[DEBUG] 步骤1: PDF文档，无需转换")
            pdf_path = file_path  # PDF文档本身就是PDF
        else:
            print(f"[DEBUG] 不支持的文档类型: {file_type}")
            return await self._parse_word_simple(file_path)
        
        # 步骤2: 文本和坐标提取
        print("[DEBUG] 步骤2: 提取文本和坐标信息")
        document_data = await self._extract_text_and_coordinates(file_path)
        
        # 步骤3: 构建字符序列映射
        print("[DEBUG] 步骤3: 构建字符序列映射")
        from app.utils.coordinate_mapper import CoordinateMapper
        mapper = CoordinateMapper()
        char_sequence_map = mapper.build_char_sequence_map(document_data)
        
        # 将映射信息添加到文档数据中
        document_data["char_sequence_map"] = char_sequence_map
        
        # 添加PDF路径信息
        document_data["pdf_path"] = pdf_path
        
        print(f"[DEBUG] ===== 文档解析完成 =====")
        print(f"[DEBUG] PDF路径: {pdf_path}")
        return document_data
    
    async def _extract_text_and_coordinates(self, pdf_path: str) -> Dict:
        """使用PyMuPDF提取文本和坐标信息 - 基于PyMuPDF示例"""
        print(f"[DEBUG] 开始解析PDF: {pdf_path}")
        
        try:
            doc = fitz.open(pdf_path)
            pages_data = []
            full_text = ""

            print(f"[DEBUG] PDF页数: {len(doc)}")

            for page_num in range(len(doc)):
                page = doc[page_num]
                print(f"[DEBUG] 解析第{page_num + 1}页")

                # 获取结构化文本信息（dict 模式包含坐标）
                text_dict = page.get_text("dict")
                page_data = {
                    "page_index": page_num,
                    "width": page.rect.width,
                    "height": page.rect.height,
                    "blocks": [],
                    "char_sequence": []  # 字符序列，用于坐标映射
                }

                print(f"[DEBUG] 页面尺寸: {page.rect.width} x {page.rect.height}")
                print(f"[DEBUG] 文本块数量: {len(text_dict.get('blocks', []))}")

                char_index = 0

                for block in text_dict["blocks"]:
                    if "lines" not in block:  # 过滤非文字块（可能是图片等）
                        print(f"[DEBUG] 跳过非文字块: {block.get('bbox', 'unknown')}")
                        continue

                    block_data = {
                        "block_index": len(page_data["blocks"]),
                        "lines": []
                    }

                    for line in block["lines"]:
                        line_data = {
                            "line_index": len(block_data["lines"]),
                            "bbox": line["bbox"],  # [x0, y0, x1, y1]
                            "spans": []
                        }

                        for span in line["spans"]:
                            span_text = span["text"]
                            if not span_text.strip():  # 跳过空文本
                                continue
                                
                            print(f"[DEBUG] 处理文本片段: '{span_text}'")

                            span_data = {
                                "text": span_text,
                                "bbox": span["bbox"],
                                "font": span["font"],
                                "size": span["size"],
                                "flags": span["flags"],
                                "color": span["color"],
                                "char_start_index": char_index,
                                "char_end_index": char_index + len(span_text)
                            }

                            # 为每个字符创建精确的坐标信息
                            char_bboxes = self._calculate_char_bboxes_precise(
                                span_text,
                                span["bbox"],
                                span["size"]
                            )

                            span_data["char_bboxes"] = char_bboxes
                            line_data["spans"].append(span_data)

                            # 更新字符序列
                            for i, char in enumerate(span_text):
                                page_data["char_sequence"].append({
                                    "char": char,
                                    "char_index": char_index + i,
                                    "bbox": char_bboxes[i] if i < len(char_bboxes) else span["bbox"],
                                    "font": span["font"],
                                    "size": span["size"],
                                    "color": span["color"]
                                })

                            char_index += len(span_text)
                            full_text += span_text

                        if line_data["spans"]:  # 只添加有内容的行
                            block_data["lines"].append(line_data)

                    if block_data["lines"]:  # 只添加有内容的块
                        page_data["blocks"].append(block_data)

                pages_data.append(page_data)

            doc.close()

            result = {
                "pages": pages_data,
                "full_text": full_text
            }

            print(f"[DEBUG] PDF解析完成，提取文本长度: {len(full_text)}")
            print(f"[DEBUG] 提取的文本内容: '{full_text[:200]}...'")
            print(f"[DEBUG] 页面数量: {len(pages_data)}")

            return result
            
        except Exception as e:
            print(f"[DEBUG] PDF解析失败: {e}")
            # 返回空结构
            return {
                "pages": [{
                    "page_index": 0,
                    "width": 612,
                    "height": 792,
                    "blocks": [],
                    "char_sequence": []
                }],
                "full_text": ""
            }
    
    async def _convert_docx_to_pdf(self, docx_path: str) -> str:
        """将Word文档转换为PDF"""
        try:
            from app.utils.format_converter import FormatConverter
            converter = FormatConverter()
            pdf_path = await converter.convert_docx_to_pdf(docx_path)
            return pdf_path
        except Exception as e:
            print(f"[DEBUG] Word转PDF失败: {e}")
            return None
    
    async def _parse_word_simple(self, file_path: str) -> Dict:
        """简化的Word文档解析（demo版本）"""
        print(f"[DEBUG] 使用简化解析处理: {file_path}")
        
        try:
            # 尝试使用python-docx解析Word文档
            from docx import Document
            doc = Document(file_path)
            
            full_text = ""
            pages_data = []
            
            # 提取所有段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text += cell.text + " "
                    full_text += "\n"
            
            # 创建页面数据结构
            page_data = {
                "page_index": 0,
                "width": 612,
                "height": 792,
                "blocks": [],
                "char_sequence": []
            }
            
            # 为每个字符创建坐标信息（模拟）
            char_index = 0
            for char in full_text:
                if char.strip():  # 只处理非空白字符
                    page_data["char_sequence"].append({
                        "char": char,
                        "char_index": char_index,
                        "bbox": [100 + (char_index % 50) * 12, 100 + (char_index // 50) * 16, 
                                112 + (char_index % 50) * 12, 116 + (char_index // 50) * 16],
                        "font": "Arial",
                        "size": 12,
                        "color": 0
                    })
                char_index += 1
            
            pages_data.append(page_data)
            
            print(f"[DEBUG] Word文档解析完成，提取文本长度: {len(full_text)}")
            print(f"[DEBUG] 提取的文本内容: '{full_text[:200]}...'")
            
            return {
                "pages": pages_data,
                "full_text": full_text.strip()
            }
            
        except Exception as e:
            print(f"[DEBUG] Word文档解析失败: {e}")
            # 返回模拟数据
            return {
                "pages": [{
                    "page_index": 0,
                    "width": 612,
                    "height": 792,
                    "blocks": [],
                    "char_sequence": []
                }],
                "full_text": "Word文档内容（解析失败）"
            }
    
    def _calculate_char_bboxes_precise(self, text: str, span_bbox: List[float], font_size: float) -> List[List[float]]:
        """计算每个字符的精确边界框 - 基于PyMuPDF示例"""
        if not text:
            return []

        x0, y0, x1, y1 = span_bbox
        char_bboxes = []

        # 使用更精确的字符宽度计算
        total_width = x1 - x0
        char_width = total_width / len(text) if len(text) > 0 else 0

        for i, char in enumerate(text):
            char_x0 = x0 + i * char_width
            char_x1 = x0 + (i + 1) * char_width
            
            # 确保坐标在合理范围内
            char_bbox = [
                max(0, char_x0),
                max(0, y0),
                min(x1, char_x1),
                min(y1, y0 + font_size)
            ]
            char_bboxes.append(char_bbox)

        return char_bboxes
    
    def _calculate_char_bboxes(self, text: str, span_bbox: List[float], font_size: float) -> List[List[float]]:
        """计算每个字符的边界框 - 兼容性方法"""
        return self._calculate_char_bboxes_precise(text, span_bbox, font_size)
